"""Generate Cover Letter for Indian Journal of Medical Microbiology"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

BASE_DIR = Path(__file__).parent.parent

def create_cover_letter():
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    p = doc.add_paragraph()
    p.add_run('December 31, 2024')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_paragraph()
    doc.add_paragraph('To,')
    doc.add_paragraph('The Editor-in-Chief')
    doc.add_paragraph('Indian Journal of Medical Microbiology')
    doc.add_paragraph('Publication of the Indian Association of Medical Microbiologists')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Subject: ').bold = True
    p.add_run('Submission of Original Research Manuscript')
    
    doc.add_paragraph()
    doc.add_paragraph('Dear Editor,')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('We submit our manuscript entitled ')
    run = p.add_run('Host-Directed Therapy Targets in Typhoid Fever: An Integrated Multi-omics Pipeline for Addressing Multidrug-Resistant Salmonella Typhi Through Autophagy Enhancement and Macrophage Reprogramming')
    run.italic = True
    p.add_run(' for publication in IJMM.')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('SUMMARY').bold = True
    
    doc.add_paragraph('Typhoid fever causes 14 million cases and 135,000 deaths annually. India carries 60% of the global burden. The emergence of XDR S. Typhi resistant to first-line agents, fluoroquinolones, AND ceftriaxone represents a critical crisis with only azithromycin remaining.')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('KEY INNOVATION').bold = True
    
    doc.add_paragraph('Host-directed therapy BYPASSES bacterial resistance by targeting host macrophage autophagy—bacteria cannot develop resistance to being degraded by host machinery.')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('PRIORITY CANDIDATES').bold = True
    
    doc.add_paragraph('- Rapamycin/Everolimus: mTOR inhibition, autophagy enhancement')
    doc.add_paragraph('- Metformin: AMPK activation, low cost, widely available')
    doc.add_paragraph('- Deferasirox: Iron chelation, nutrient deprivation')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('DECLARATIONS').bold = True
    
    doc.add_paragraph('- Original work, not under consideration elsewhere')
    doc.add_paragraph('- Computational study, no ethics approval required')
    doc.add_paragraph('- No conflicts of interest')
    
    doc.add_paragraph()
    doc.add_paragraph('This work addresses a critical AMR challenge highly relevant to Indian medical microbiology.')
    
    doc.add_paragraph()
    doc.add_paragraph('Respectfully submitted,')
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('Dr. Siddalingaiah H S').bold = True
    doc.add_paragraph('Professor, Community Medicine')
    doc.add_paragraph('SIMS Tumkur, Karnataka')
    doc.add_paragraph('Email: hssling@yahoo.com')
    doc.add_paragraph('ORCID: 0000-0002-4771-8285')
    
    doc.save(str(BASE_DIR / 'manuscripts' / 'CoverLetter_IJMM.docx'))
    print('Created: CoverLetter_IJMM.docx')

if __name__ == '__main__':
    create_cover_letter()
