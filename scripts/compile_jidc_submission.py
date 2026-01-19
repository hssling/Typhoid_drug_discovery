"""
Typhoid HDT v4.1: JIDC Submission Package Compiler
Compiles all audited assets and generates a specialized Cover Letter with APC Waiver.
"""

import shutil
import pandas as pd
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = Path(__file__).parent.parent
SUBMISSION_DIR = BASE_DIR / 'submission' / 'jidc'

def compile_assets():
    print("="*60)
    print("COMPILING JIDC SUBMISSION PACKAGE (APC-FREE STRATEGY)")
    print("="*60)
    
    # 1. Copy Figures & Tables
    print("Copying Figures & Tables...")
    fig_src = BASE_DIR / 'outputs' / 'figures'
    table_src = BASE_DIR / 'outputs' / 'tables' / 'targets_ranked_v5_global.csv'
    
    if fig_src.exists():
        for fig in fig_src.glob('*.png'):
            shutil.copy(fig, SUBMISSION_DIR / 'figures' / fig.name)
            
    if table_src.exists():
        shutil.copy(table_src, SUBMISSION_DIR / 'tables' / 'Table1_Global_Priorities.csv')

    # 2. Copy Audit Documents (Proof of Authenticity)
    print("Copying Audit Documents...")
    shutil.copy(BASE_DIR / 'docs' / 'SCIENTIFIC_VERIFICATION_REPORT.md', SUBMISSION_DIR / 'docs' / 'Audit_Authenticity_Report.md')
    shutil.copy(BASE_DIR / 'docs' / 'v4_peer_review_report.md', SUBMISSION_DIR / 'docs' / 'Peer_Review_Response_Log.md')
    shutil.copy(BASE_DIR / 'docs' / 'v4_validation_protocol.md', SUBMISSION_DIR / 'docs' / 'Validation_Protocol.md')

    # 3. Generate JIDC Cover Letter + Waiver Request
    print("Generating Cover Letter & Waiver...")
    doc = Document()
    
    title = doc.add_paragraph()
    run = title.add_run("SUBMISSION TO: JOURNAL OF INFECTION IN DEVELOPING COUNTRIES")
    run.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("\nDear Editorial Board,\n")
    doc.add_paragraph(
        "Attached is our manuscript, \"Convergent Evidence-Based Discovery of Broad-Spectrum Host-Directed Therapies for Enteric Fever,\" for consideration in JIDC."
    )
    doc.add_paragraph(
        "This study is highly relevant to JIDC as it addresses the crisis of MDR/XDR Typhoid in low-and-middle-income countries (LMICs). We utilize an authentic multi-omics approach, causally de-risked by human GWAS data, to identify host targets that are resilient to antimicrobial resistance."
    )
    
    doc.add_paragraph("\n[APC WAIVER REQUEST]\n", style='Heading 1')
    doc.add_paragraph(
        "As this research focuses primarily on addressing critical health disparities in developing nations and was conducted as an open-science initiative to support public health in enteric fever endemic regions, we respectfully request a full waiver of the Article Processing Charge (APC). Our goal is to ensure this discovery remains accessible to the researchers and clinicians in the countries most burdened by this disease."
    )
    
    doc.add_paragraph("\nSincerely,\n\nDr. Siddalingaiah H S\nLead Researcher")
    
    doc.save(SUBMISSION_DIR / 'docs' / 'Cover_Letter_and_Waiver.docx')

    # 4. Generate Final Hardened Manuscript (v4.1)
    # We will copy the existing submission package but rename it clearly
    shutil.copy(BASE_DIR / 'manuscripts' / 'Typhoid_HDT_v4_SUBMISSION_PACKAGE.docx', SUBMISSION_DIR / 'docs' / 'Main_Manuscript_v4.1_Hardened.docx')

    print("\n" + "="*60)
    print("Asset Compilation Complete.")
    print(f"Location: {SUBMISSION_DIR.relative_to(BASE_DIR)}")
    print("="*60)

if __name__ == "__main__":
    compile_assets()
